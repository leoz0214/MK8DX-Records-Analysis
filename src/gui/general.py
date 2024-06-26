"""General analysis for all courses: 150cc, 200cc, both 150cc and 200cc."""
import datetime as dt
import enum
import tkinter as tk
from dataclasses import dataclass
from tkinter import filedialog
from tkinter import messagebox
from tkinter import ttk

import docx

import main
from gutils import (
    lato, TablesFrame, RankTable, get_date_range_string,
    COURSE_COL_WIDTH, CC_COL_WIDTH, IMPROVEMENT_COL_WIDTH, COUNT_COL_WIDTH,
    DAYS_BY_PLAYER_COLS, RECORDS_BY_PLAYER_COLS, DAYS_BY_COUNTRY_COLS,
    RECORDS_BY_COUNTRY_COLS, RECORDS_BY_CHARACTER_COLS, RECORDS_BY_KART_COLS,
    RECORDS_BY_TYRES_COLS, RECORDS_BY_GLIDER_COLS,
    UpdateDateRangeToplevel, TableExportationFrame, get_raw_records_columns,
    TableDocumentExportationOptions, add_doctable, save_docx_as_pdf)
from utils import (
    get_150cc_records, get_200cc_records, get_uniques,
    get_most_recent_snapshot_date_time)


class ExportTable(enum.Enum):
    # Various tables that can be exported as CSV (one) or XLSX (sheets).
    raw_records = "Raw Records"
    days_by_player = "Days by Player"
    records_by_player = "Records by Player"
    days_by_country = "Days by Country"
    records_by_country = "Records by Country"
    records_by_character = "Records by Character"
    records_by_kart = "Records by Kart"
    records_by_tyres = "Records by Tyres"
    records_by_glider = "Records by Glider"
    current_raw_records = "Current Raw Records"
    current_days_by_player = "Current Days by Player"
    current_records_by_player = "Current Records by Player"
    current_days_by_country = "Current Days by Country"
    current_records_by_country = "Current Records by Country"
    current_records_by_character = "Current Records by Character"
    current_records_by_kart = "Current Records by Kart"
    current_records_by_tyres = "Current Records by Tyres"
    current_records_by_glider = "Current Records by Glider"
    biggest_course_cc_improvements = "Biggest Course-CC Improvements"


@dataclass
class BasicDocumentOptions:
    """Include the summary stats / current summary stats or not."""
    summary_stats: bool
    current_summary_stats: bool


DOCUMENT_GENERAL_EXPORT_TABLES = (
    ExportTable.days_by_player,
    ExportTable.records_by_player, ExportTable.days_by_country,
    ExportTable.records_by_country, ExportTable.records_by_character,
    ExportTable.records_by_kart, ExportTable.records_by_tyres,
    ExportTable.records_by_glider,
    ExportTable.current_days_by_player, ExportTable.current_records_by_player,
    ExportTable.current_days_by_country,
    ExportTable.current_records_by_country,
    ExportTable.current_records_by_character,
    ExportTable.current_records_by_kart, ExportTable.current_records_by_tyres,
    ExportTable.current_records_by_glider,
    ExportTable.biggest_course_cc_improvements
)
# Just the document export tables but with raw records and current raw records.
GENERAL_EXPORT_TABLES = (
    ExportTable.raw_records, *DOCUMENT_GENERAL_EXPORT_TABLES[
        :DOCUMENT_GENERAL_EXPORT_TABLES.index(
            ExportTable.current_days_by_player)]
) + (
    ExportTable.current_raw_records, *DOCUMENT_GENERAL_EXPORT_TABLES[
        DOCUMENT_GENERAL_EXPORT_TABLES.index(
            ExportTable.current_days_by_player):])
EXPORT_TABLES_COLUMNS = {
    ExportTable.days_by_player: DAYS_BY_PLAYER_COLS,
    ExportTable.records_by_player: RECORDS_BY_PLAYER_COLS,
    ExportTable.days_by_country: DAYS_BY_COUNTRY_COLS,
    ExportTable.records_by_country: RECORDS_BY_COUNTRY_COLS,
    ExportTable.records_by_character: RECORDS_BY_CHARACTER_COLS,
    ExportTable.records_by_kart: RECORDS_BY_KART_COLS,
    ExportTable.records_by_tyres: RECORDS_BY_TYRES_COLS,
    ExportTable.records_by_glider: RECORDS_BY_GLIDER_COLS
}
# All current variations have the same columns, make it so.
for table, columns in EXPORT_TABLES_COLUMNS.copy().items():
    name = str(table).split(".")[1]
    EXPORT_TABLES_COLUMNS[getattr(ExportTable, f"current_{name}")] = columns


COURSE_CC_IMPROVEMENT_COLS = ("Course", "CC", "Count", "Improvement")


def get_columns_and_records_by_table(
    analysis_frame: "GeneralAnalysisFrame", table: ExportTable
) -> tuple[tuple[str], list[tuple]]:
    """Returns records given analysis frame and desired table."""
    if table == ExportTable.biggest_course_cc_improvements:
        records = analysis_frame.course_cc_improvements_records
        return COURSE_CC_IMPROVEMENT_COLS, records
    columns = EXPORT_TABLES_COLUMNS[table]
    attribute = str(table).split(".")[1]
    # If it is a current variant, use the current records data
    # and dummy change the current attribute to normal attribute.
    if not attribute.startswith("current_"):
        tables = analysis_frame.overall_frame.tables_frame
    else:
        tables = analysis_frame.current_frame.tables_frame
        attribute = attribute.removeprefix("current_")
        table = getattr(ExportTable, attribute)
    records = {
        ExportTable.days_by_player: tables.days_by_player_records,
        ExportTable.records_by_player: tables.player_record_counts,
        ExportTable.days_by_country: tables.days_by_country_records,
        ExportTable.records_by_country: tables.country_record_counts,
        ExportTable.records_by_character: tables.character_record_counts,
        ExportTable.records_by_kart: tables.kart_record_counts,
        ExportTable.records_by_tyres: tables.tyres_record_counts,
        ExportTable.records_by_glider: tables.glider_record_counts
    }[table]
    return columns, records


class GeneralAnalysisFrame(tk.Frame):
    """General analysis for all courses."""

    def __init__(
        self, tab: "main.Tab", mode: "main.GeneralMode",
        min_date: dt.date | None, max_date: dt.date | None
    ) -> None:
        super().__init__(tab)
        self.tab = tab
        self.mode = mode
        self.mode_string = {
            main.GeneralMode.only150.value: "150cc",
            main.GeneralMode.only200.value: "200cc",
            main.GeneralMode.both.value: "150cc and 200cc"
        }[self.mode.value]
        if mode.value == main.GeneralMode.only150.value:
            self.records = get_150cc_records(min_date, max_date)
        elif mode.value == main.GeneralMode.only200.value:
            self.records = get_200cc_records(min_date, max_date)
        else:
            self.records = (
                get_150cc_records(min_date, max_date)
                + get_200cc_records(min_date, max_date))
        self.min_date = min_date
        self.max_date = max_date
        self.title = tk.Label(
            self, font=lato(25, True),
            text=f"General {self.mode_string} Analysis")
        # Two variants of tables - tables for overall, and tables
        # for current records only.
        self.notebook = ttk.Notebook(self)
        self.overall_frame = OverallRecordsDataFrame(self)
        self.current_frame = CurrentRecordsDataFrame(self)
        self.notebook.add(self.overall_frame, text="Overall")
        self.notebook.add(self.current_frame, text="Current")
        # Additional table - biggest improvements by course/CC.
        # Using 2-lists (slowest, fastest)
        course_cc_worst_best_times = {}
        course_cc_counter = {}
        for record in self.records:
            key = (record.course, record.is200)
            if key not in course_cc_worst_best_times:
                course_cc_worst_best_times[key] = [record.time, record.time]
                course_cc_counter[key] = 1
                continue
            if record.time > course_cc_worst_best_times[key][0]:
                course_cc_worst_best_times[key][0] = record.time
            if record.time < course_cc_worst_best_times[key][1]:
                course_cc_worst_best_times[key][1] = record.time
            course_cc_counter[key] += 1
        course_cc_improvements = {
            course_cc: slowest - fastest
            for course_cc,
                (slowest, fastest) in course_cc_worst_best_times.items()}
        course_cc_improvements = dict(
            sorted(course_cc_improvements.items(),
                key=lambda item: item[1], reverse=True))
        self.course_cc_improvements_records = []
        for course_cc, improvement in course_cc_improvements.items():
            record = (
                course_cc[0], 200 if course_cc[1] else 150,
                course_cc_counter[course_cc], round(improvement / 1000, 3))
            self.course_cc_improvements_records.append(record)
        self.course_cc_improvements_table = RankTable(
            self, "Biggest Course/CC Improvements",
            COURSE_CC_IMPROVEMENT_COLS, self.course_cc_improvements_records,
            column_widths=(
                COURSE_COL_WIDTH, CC_COL_WIDTH, COUNT_COL_WIDTH,
                IMPROVEMENT_COL_WIDTH))
        self.options_frame = OptionsFrame(self)

        self.title.pack(pady=5)
        self.notebook.pack(pady=5)
        self.course_cc_improvements_table.pack(pady=5)
        self.options_frame.pack(pady=5)

    def export(self) -> None:
        """Allows exportation of the data to various file formats."""
        ExportationToplevel(self)

    def update_date_range(
        self, min_date: dt.date | None, max_date: dt.date | None
    ) -> None:
        """Updates the date range by refreshing the window."""
        # Dummy change, then call refresh data.
        self.min_date = min_date
        self.max_date = max_date
        self.refresh_data()

    def refresh_data(self) -> None:
        """Reset the screen in case of a new snapshot."""
        self.destroy()
        GeneralAnalysisFrame(
            self.tab, self.mode, self.min_date, self.max_date).pack()


class RecordsDataFrame(tk.Frame):
    """Parent records data frame for both current/overall records data."""

    def __init__(self, master: GeneralAnalysisFrame) -> None:
        super().__init__(master)
        (
            self.unique_players, self.unique_countries, unique_characters,
            unique_karts, unique_tyres, unique_gliders
        ) = get_uniques(self.records)
        self.info_label = tk.Label(
            self, font=lato(15, italic=True), text=(
                f"Records: {len(self.records)} | "
                f"Players: {len(self.unique_players)} | "
                f"Countries: {len(self.unique_countries)} | "
                f"Characters: {len(unique_characters)} | "
                f"Karts: {len(unique_karts)} | "
                f"Tyres: {len(unique_tyres)} | "
                f"Gliders: {len(unique_gliders)}"))
        self.tables_frame = TablesFrame(self)
        self.info_label.pack(pady=5)
        self.tables_frame.pack(pady=5)


class OverallRecordsDataFrame(RecordsDataFrame):
    """
    Operates on all the records data for a given CC in the date range,
    providing relevant info and tables.
    """

    def __init__(self, master: GeneralAnalysisFrame) -> None:
        self.records = master.records
        super().__init__(master)


class CurrentRecordsDataFrame(RecordsDataFrame):
    """
    Obtains the current records for each course/CC
    and provides analysis on them.
    """

    def __init__(self, master: GeneralAnalysisFrame) -> None:
        records = master.records
        by_course_cc = {}
        for record in records:
            key = (record.course, record.is200)
            # Ensure ties are handled correctly and considered current.
            if key not in by_course_cc:
                by_course_cc[key] = [record]
                continue
            if record.time == by_course_cc[key][0].time:
                by_course_cc[key].append(record)
            elif record.time < by_course_cc[key][0].time:
                by_course_cc[key] = [record]
        self.records = []
        for records in by_course_cc.values():
            self.records.extend(records)
        super().__init__(master)


class OptionsFrame(tk.Frame):
    """Various buttons and info to interact with analysis tab."""

    def __init__(self, master: GeneralAnalysisFrame) -> None:
        super().__init__(master)
        self.export_button = ttk.Button(
            self, text="Export", command=master.export)
        self.close_button = ttk.Button(
            self, text="Close", command=master.tab.close)
        date_range = get_date_range_string(master.min_date, master.max_date)
        self.date_range_label = tk.Label(
            self, text=f"Date Range: {date_range}", width=30)
        self.update_date_range_button = ttk.Button(
            self, text="Update",
            command=lambda: UpdateDateRangeToplevel(
                master.update_date_range, master.min_date, master.max_date))
        data_time = get_most_recent_snapshot_date_time().replace(microsecond=0)
        self.data_time_label = tk.Label(
            self, text=f"Data collected at: {data_time}", width=30)
        self.refresh_data_button = ttk.Button(
            self, text="Refresh", command=master.refresh_data)
        self.export_button.grid(row=0, column=0, columnspan=2, padx=5, pady=5)
        self.close_button.grid(row=0, column=2, columnspan=2, padx=5, pady=5)
        self.date_range_label.grid(row=1, column=0, padx=5, pady=5)
        self.update_date_range_button.grid(row=1, column=1, padx=5, pady=5)
        self.data_time_label.grid(row=1, column=2, padx=5, pady=5)
        self.refresh_data_button.grid(row=1, column=3, padx=5, pady=5)


class ExportationToplevel(tk.Toplevel):
    """Exportation to CSV, XLSX, DOCX and PDF."""

    def __init__(self, master: GeneralAnalysisFrame) -> None:
        super().__init__()
        self.master: GeneralAnalysisFrame = master
        self.title(f"{master.title.cget('text')} - Data Exportation")
        self.grab_set()
        self.notebook = ttk.Notebook(self)
        self.table_exportation_frame = GeneralTableExportationFrame(self)
        self.document_exportation_frame = DocumentExportationFrame(self)
        self.notebook.add(self.table_exportation_frame, text="CSV/XLSX")
        self.notebook.add(self.document_exportation_frame, text="DOCX/PDF")
        self.notebook.pack()


class GeneralTableExportationFrame(TableExportationFrame):
    """Genera analysis table exportation subclass."""

    def __init__(self, master: ExportationToplevel) -> None:
        super().__init__(master, GENERAL_EXPORT_TABLES)

    def get_columns_and_records(
        self, table: ExportTable
    ) -> tuple[tuple[str], list[tuple]]:
        """Gets columns and records for a general table."""
        if table in (ExportTable.raw_records, ExportTable.current_raw_records):
            if table == ExportTable.raw_records:
                db_records = self.master.master.records
            else:
                db_records = self.master.master.current_frame.records
            # Assume baby park included in general analysis - 7 laps max.
            columns = get_raw_records_columns(7)
            records = []
            for record in db_records:
                export_record = [
                    record.course, record.is200, record.date, record.time,
                    record.player, record.country, record.days]
                for field in ("lap_times", "coins", "mushrooms"):
                    value = getattr(record, field)
                    if value is not None:
                        export_record.extend(value)
                        export_record.extend([None] * (7 - len(value)))
                    else:
                        export_record.extend([None] * 7)
                export_record.extend(
                    (record.character, record.kart, record.tyres,
                        record.glider, record.video_link))
                records.append(export_record)
            return columns, records
        return get_columns_and_records_by_table(self.master.master, table)  


class DocumentExportationFrame(tk.Frame):
    """Allows the user to export to DOCX or PDF."""

    def __init__(self, master: ExportationToplevel) -> None:
        super().__init__(master)
        self.title = tk.Label(
            self, font=lato(25, True), text="Export Document")
        self.info_label = tk.Label(
            self, text="Select the components below to generate.")
        self.basic_label = tk.Label(self, font=lato(15, True), text="Basic")
        self.basic_frame = BasicDocumentExportationOptions(self)
        self.tables_label = tk.Label(self, font=lato(15, True), text="Tables")
        self.tables_frame = TableDocumentExportationOptions(
            self, DOCUMENT_GENERAL_EXPORT_TABLES)
        self.export_docx_button = ttk.Button(
            self, text="Export DOCX", command=self.export_docx)
        self.export_pdf_button = ttk.Button(
            self, text="Export PDF [WARNING: Unstable!]",
            command=self.export_pdf)
        
        self.title.pack(pady=5)
        self.info_label.pack(pady=5)
        self.basic_label.pack(pady=5)
        self.basic_frame.pack(pady=5)
        self.tables_label.pack(pady=5)
        self.tables_frame.pack(pady=5)
        self.export_docx_button.pack(pady=5)
        self.export_pdf_button.pack(pady=5)
    
    def _generate_docx(self) -> docx.document.Document:
        # Generate DOCX from the configuration options.
        basic_options = self.basic_frame.options
        tables = self.tables_frame.selected
        rows_per_table = self.tables_frame.max_count
        document: docx.document.Document = docx.Document()
        document.add_heading(self.master.master.title.cget("text"), 0)
        if basic_options.summary_stats or basic_options.current_summary_stats:
            self._add_basic(document, basic_options)
        if tables:
            self._add_tables(document, tables, rows_per_table)
        document.add_heading("Data Context", 1)
        options_frame = self.master.master.options_frame
        date_range_text = options_frame.date_range_label.cget("text")
        data_time_text = options_frame.data_time_label.cget("text")
        document.add_paragraph(date_range_text)
        document.add_paragraph(data_time_text)
        return document

    def _add_basic(
        self, document: docx.document.Document,
        basic_options: BasicDocumentOptions
    ) -> None:
        for var, heading, frame in (
            (basic_options.summary_stats, "Summary Stats",
                self.master.master.overall_frame),
            (basic_options.current_summary_stats, "Current Summary Stats",
                self.master.master.current_frame)
        ):
            if var:
                document.add_heading(heading, 1)
                info_label = frame.info_label
                points = map(str.strip, info_label.cget("text").split("|"))
                for point in points:
                    document.add_paragraph(point, style="List Bullet")
    
    def _add_tables(
        self, document: docx.document.Document,
        tables: list[ExportTable], rows_per_table: int
    ) -> None:
        document.add_heading("Tables", 1)
        for table in tables:
            document.add_heading(table.value, 2)
            columns, records = get_columns_and_records_by_table(
                self.master.master, table)
            columns = ("#", *columns)
            records = records[:rows_per_table]
            add_doctable(document, columns, records)

    def export_docx(self) -> None:
        """Exports to DOCX."""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx", filetypes=(("DOCX", ".docx"),),
            title="Save DOCX", parent=self)
        if not file_path:
            return
        try:
            document = self._generate_docx()
            document.save(file_path)
        except Exception as e:
            messagebox.showerror(
                "Error",
                    "Unfortunately, an error occurred "
                    f"while exporting to DOCX: {e}", parent=self)
    
    def export_pdf(self) -> None:
        """Exports to PDF (by generating DOCX and then converting to PDF)."""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".pdf", filetypes=(("PDF", ".pdf"),),
            title="Save PDF", parent=self)
        if not file_path:
            return
        try:
            document = self._generate_docx()
            save_docx_as_pdf(document, file_path)
        except Exception as e:
            messagebox.showerror(
                "Error",
                    "Unfortunately, an error occurred "
                    f"while exporting to PDF: {e}", parent=self)


class BasicDocumentExportationOptions(tk.Frame):
    """
    Sets whether or not to include summary stats / current summary stats
    in the document.
    """

    def __init__(self, master: DocumentExportationFrame) -> None:
        super().__init__(master)
        self._summary_stats = tk.BooleanVar(value=True)
        self._current_summary_stats = tk.BooleanVar(value=True)
        self.summary_stats_checkbutton = ttk.Checkbutton(
            self, text="Summary Stats", variable=self._summary_stats)
        self.current_summary_stats_checkbutton = ttk.Checkbutton(
            self, text="Current Summary Stats",
            variable=self._current_summary_stats)
        self.summary_stats_checkbutton.pack()
        self.current_summary_stats_checkbutton.pack()
    
    @property
    def options(self) -> BasicDocumentOptions:
        return BasicDocumentOptions(
            self._summary_stats.get(), self._current_summary_stats.get())
