"""Data analysis for course/CC."""
import datetime as dt
import enum
import io
import platform
import tkinter as tk
from dataclasses import dataclass
from tkinter import filedialog
from tkinter import messagebox
from tkinter import ttk
from typing import Callable

import docx
import docx.document
import matplotlib.dates as mdates
from docx.shared import Inches
from matplotlib import pyplot as plt

import main
from gutils import (
    lato, TablesFrame, MillisecondsFormatter, get_date_range_string,
    UpdateDateRangeToplevel, get_raw_records_columns,
    DAYS_BY_PLAYER_COLS, RECORDS_BY_PLAYER_COLS, DAYS_BY_COUNTRY_COLS,
    RECORDS_BY_COUNTRY_COLS, RECORDS_BY_CHARACTER_COLS,
    RECORDS_BY_KART_COLS, RECORDS_BY_TYRES_COLS,
    RECORDS_BY_GLIDER_COLS, TableExportationFrame,
    TableDocumentExportationOptions, add_doctable, save_docx_as_pdf)
from utils import (
    get_course_cc_records, ms_to_finish_time, get_lap_count,
    get_most_recent_snapshot_date_time, get_uniques)


GRAPH_DATES_INTERVALS = 8
GRAPH_IMAGE_DPI = 75
GRAPH_TIME_FORMAT = f"%{'#' if platform.system() == 'Windows' else '-'}M:%S.%f"
# Document export graph size
GRAPH_WIDTH = Inches(5)
GRAPH_HEIGHT = Inches(3.75)


class ExportTable(enum.Enum):
    """
    Various tables that can be exported as CSV (one) or XLSX (sheets).
    Also borrowed in the export document table section.
    """
    raw_records = "Raw Records"
    days_by_player = "Days by Player"
    records_by_player = "Records by Player"
    days_by_country = "Days by Country"
    records_by_country = "Records by Country"
    records_by_character = "Records by Character"
    records_by_kart = "Records by Kart"
    records_by_tyres = "Records by Tyres"
    records_by_glider = "Records by Glider"


class ExportGraph(enum.Enum):
    """Graphs that can be exported as part of the document."""
    times = "Time against Date"
    lap_times = "Lap Times against Date"
    coins = "Coins against Date"
    mushrooms = "Mushrooms against Date"


@dataclass
class BasicDocumentOptions:
    """Basic document options: generate current records and summary stats."""
    current_records: bool
    summary_stats: bool


DOCUMENT_COURSE_CC_EXPORT_TABLES = (
    ExportTable.days_by_player,
    ExportTable.records_by_player, ExportTable.days_by_country,
    ExportTable.records_by_country, ExportTable.records_by_character,
    ExportTable.records_by_kart, ExportTable.records_by_tyres,
    ExportTable.records_by_glider
)
COURSE_CC_EXPORT_TABLES = (
    ExportTable.raw_records,) + DOCUMENT_COURSE_CC_EXPORT_TABLES
EXPORT_TABLES_COLUMNS = {
    ExportTable.days_by_player: DAYS_BY_PLAYER_COLS,
    ExportTable.records_by_player: RECORDS_BY_PLAYER_COLS,
    ExportTable.days_by_country: DAYS_BY_COUNTRY_COLS,
    ExportTable.records_by_country: RECORDS_BY_COUNTRY_COLS,
    ExportTable.records_by_character: RECORDS_BY_CHARACTER_COLS,
    ExportTable.records_by_kart: RECORDS_BY_KART_COLS,
    ExportTable.records_by_tyres: RECORDS_BY_TYRES_COLS,
    ExportTable.records_by_glider: RECORDS_BY_GLIDER_COLS
}

COURSE_CC_EXPORT_GRAPHS = (
    ExportGraph.times, ExportGraph.lap_times,
    ExportGraph.coins, ExportGraph.mushrooms
)


def set_up_dates_xaxis(dates: list[dt.date]) -> None:
    """Sets the date x-axis, given a list of dates."""
    plt.gca().xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m-%d"))
    plt.gca().xaxis.set_major_locator(mdates.DayLocator())
    days_range = (dates[-1] - dates[0]).days
    plt.gca().xaxis.set_major_locator(
        mdates.DayLocator(
            interval=max(days_range // GRAPH_DATES_INTERVALS, 1)))
    plt.gcf().autofmt_xdate()


def get_records_from_tables(
    tables: "TablesFrame", table: ExportTable
) -> list[tuple]:
    """Returns the records given the tables frame and export table."""
    return {
        ExportTable.days_by_player: tables.days_by_player_records,
        ExportTable.records_by_player: tables.player_record_counts,
        ExportTable.days_by_country: tables.days_by_country_records,
        ExportTable.records_by_country: tables.country_record_counts,
        ExportTable.records_by_character: tables.character_record_counts,
        ExportTable.records_by_kart: tables.kart_record_counts,
        ExportTable.records_by_tyres: tables.tyres_record_counts,
        ExportTable.records_by_glider: tables.glider_record_counts
    }[table]


class CourseCcAnalysis(tk.Frame):
    """Displays course/CC records insights and allows for data exporting."""

    def __init__(
        self, tab: main.Tab, course: str, is200: bool,
        min_date: dt.date | None, max_date: dt.date | None
    ) -> None:
        super().__init__(tab)
        self.course = course
        self.is200 = is200
        self.min_date = min_date
        self.max_date = max_date
        self.tab = tab
        self.records = get_course_cc_records(
            self.course, self.is200, self.min_date, self.max_date)
        (
            self.unique_players, self.unique_countries, unique_characters,
            unique_karts, unique_tyres, unique_gliders
        ) = get_uniques(self.records)
        self.title = tk.Label(
            self, font=lato(25, True),
            text=f"{self.course} {200 if self.is200 else 150}cc")
        self.current_records_frame = CurrentRecordsFrame(self)

        self.info_label = tk.Label(
            self, font=lato(15, italic=True), text=(
                f"Records: {len(self.records)} | "
                f"Players: {len(self.unique_players)} | "
                f"Countries: {len(self.unique_countries)} | "
                f"Characters: {len(unique_characters)} | "
                f"Karts: {len(unique_karts)} | "
                f"Tyres: {len(unique_tyres)} | "
                f"Gliders: {len(unique_gliders)}"))
        self.tables_frame = TablesFrame(self)
        self.options_frame = OptionsFrame(self)

        self.title.pack(pady=5)
        self.current_records_frame.pack(pady=5)
        self.info_label.pack(pady=5)
        self.tables_frame.pack(pady=5)
        self.options_frame.pack(pady=5)
    
    def open_graphs(self) -> None:
        """Opens the graphs toplevel to view graphs based on the stats."""
        GraphsToplevel(self)

    def export(self) -> None:
        """Allows exportation of the data to various file formats."""
        ExportationToplevel(self)

    def update_date_range(
        self, min_date: dt.date | None, max_date: dt.date | None
    ) -> None:
        """Updates the date range by refreshing the window."""
        # Dummy change, then call refresh data.
        self.min_date = min_date
        self.max_date = max_date
        self.refresh_data()

    def refresh_data(self) -> None:
        """Refreshes the window - in case new snapshot is available."""
        self.destroy()
        CourseCcAnalysis(
            self.tab, self.course, self.is200,
            self.min_date, self.max_date).pack()


class OptionsFrame(tk.Frame):
    """Various buttons and info to interact with analysis tab."""

    def __init__(self, master: CourseCcAnalysis) -> None:
        super().__init__(master)
        self.graphs_button = ttk.Button(
            self, text="Open Graphs", command=master.open_graphs)
        self.export_button = ttk.Button(
            self, text="Export", command=master.export)
        self.close_button = ttk.Button(
            self, text="Close", command=master.tab.close)
        date_range = get_date_range_string(master.min_date, master.max_date)
        self.date_range_label = tk.Label(
            self, text=f"Date Range: {date_range}", width=30)
        self.update_date_range_button = ttk.Button(
            self, text="Update",
            command=lambda: UpdateDateRangeToplevel(
                master.update_date_range, master.min_date, master.max_date))
        data_time = get_most_recent_snapshot_date_time().replace(microsecond=0)
        self.data_time_label = tk.Label(
            self, text=f"Data collected at: {data_time}", width=30)
        self.refresh_data_button = ttk.Button(
            self, text="Refresh", command=master.refresh_data)
        self.graphs_button.grid(row=0, column=0, padx=5, pady=5)
        self.export_button.grid(row=0, column=1, padx=5, pady=5)
        self.close_button.grid(row=0, column=2, padx=5, pady=5)
        self.date_range_label.grid(row=1, column=0, padx=5, pady=5)
        self.update_date_range_button.grid(row=1, column=1, padx=5, pady=5)
        self.data_time_label.grid(row=1, column=2, padx=5, pady=5)
        self.refresh_data_button.grid(row=1, column=3, padx=5, pady=5)


class CurrentRecordsFrame(tk.Frame):
    """Displays the current records (possibly ties)."""

    def __init__(self, master: CourseCcAnalysis) -> None:
        super().__init__(master)
        try:
            current_record_time = min(
                master.records, key=lambda record: record.time).time
            self.current_records = [
                record for record in master.records
                    if record.time == current_record_time]
            text = (
                f"Current record{'s' if len(self.current_records) > 1 else ''}"
            )
            self.label = tk.Label(self, font=lato(15, True), text=text)
            if len(self.current_records) > 2:
                # n-way tie where n > 2, not enough space to display, skip.
                # Really very rare... 3 way ties have happened before for sure.
                text = (
                    f"{len(self.current_records)}-way tie! "
                    "(too much to display)")
                self.label = tk.Label(self, font=lato(15, True), text=text)
                self.current_records = []
        except ValueError:
            self.current_records = []
            self.label = tk.Label(
                self, font=lato(15, True),
                text="No records within the dates!")
        self.label.pack()
        for record in self.current_records:
            finish_time = ms_to_finish_time(record.time)
            lap_times = record.lap_times
            if lap_times is not None:
                lap_times = tuple(round(lap / 1000, 3) for lap in lap_times)
            record_text = (
                f"Date: {record.date} | Time: {finish_time} | "
                f"Player: {record.player} | Country: {record.country} | "
                f"Days lasted: {record.days}\n"
                f"Lap times: {lap_times} | Coins per lap: {record.coins} | "
                f"Mushrooms per lap: {record.mushrooms}\n"
                f"Character: {record.character} | Kart: {record.kart} | "
                f"Tyres: {record.tyres} | Glider: {record.glider}")
            record_label = tk.Label(self, text=record_text)
            record_label.pack()
            if record.video_link is not None:
                video_link_entry = ttk.Entry(self, width=48)
                video_link_entry.insert(0, record.video_link)
                video_link_entry.config(state="readonly")
                video_link_entry.pack()


class GraphsToplevel(tk.Toplevel):
    """Display various graphs for a particular course/CC."""

    def __init__(self, master: CourseCcAnalysis) -> None:
        super().__init__()
        self.master: CourseCcAnalysis = master
        title = f"{master.course} {200 if master.is200 else 150}cc - Graphs"
        self.title(title)
        self.title_label = tk.Label(self, font=lato(25, True), text=title)

        self.time_against_date_graph = GraphFrame(
            self, self.plot_time_against_date)
        self.laps_against_date_graph = GraphFrame(
            self, self.plot_laps_against_date)
        self.coins_against_date_graph = GraphFrame(
            self, self.plot_coins_against_date)
        self.mushrooms_against_date_graph = GraphFrame(
            self, self.plot_mushrooms_against_date)
        
        self.title_label.grid(row=0, column=0, columnspan=2, padx=5, pady=5)
        self.time_against_date_graph.grid(row=1, column=0, padx=5, pady=5)
        self.laps_against_date_graph.grid(row=1, column=1, padx=5, pady=5)
        self.coins_against_date_graph.grid(row=2, column=0, padx=5, pady=5)
        self.mushrooms_against_date_graph.grid(row=2, column=1, padx=5, pady=5)

    def plot_time_against_date(self) -> bool:
        """
        Plots the finish time against date graph, returning True if successful.
        """
        times = []
        dates = []
        for record in self.master.records:
            if record.date is None:
                continue
            # Use dummy datetimes (only time matters).
            minutes, ms = divmod(record.time, 60000)
            seconds, ms = divmod(ms, 1000)
            times.append(
                dt.datetime(1970, 1, 1,
                    minute=minutes,second=seconds, microsecond=ms * 1000))
            dates.append(record.date)
        if not times:
            return False
        plt.clf()
        set_up_dates_xaxis(dates)
        plt.gca().yaxis.set_major_formatter(
            MillisecondsFormatter(GRAPH_TIME_FORMAT))
        plt.plot(dates, times)
        plt.title("Time against Date")
        plt.xlabel("Date")
        plt.ylabel("Time")
        return True

    def plot_laps_against_date(self) -> bool:
        """Plots the laps against date graph."""
        laps = []
        dates = []
        for record in self.master.records:
            if record.date is None or record.lap_times is None:
                continue
            for i, lap_time in enumerate(record.lap_times):
                # Again, Use dummy datetimes (only time matters).
                seconds, ms = divmod(lap_time, 1000)
                dummy_date_time = dt.datetime(
                    1970, 1, 1, second=seconds, microsecond=ms * 1000)
                if i >= len(laps):
                    laps.append([dummy_date_time])
                else:
                    laps[i].append(dummy_date_time)
            dates.append(record.date)
        if not laps:
            return False
        plt.clf()
        set_up_dates_xaxis(dates)
        plt.gca().yaxis.set_major_formatter(MillisecondsFormatter("%S.%f"))
        for n, lap_times in enumerate(laps, 1):
            plt.plot(dates, lap_times, label=f"Lap {n}")
        plt.title("Lap Times against Date")
        plt.xlabel("Date")
        plt.ylabel("Lap Time")
        plt.legend(loc="upper left")
        return True
    
    def _plot_item_against_date(self, field: str) -> bool:
        # Too similar to repeat - plot coins/mushrooms against date.
        counts = []
        dates = []
        for record in self.master.records:
            if record.date is None or getattr(record, field) is None:
                continue
            for i, lap_count in enumerate(getattr(record, field)):
                if i >= len(counts):
                    counts.append([lap_count])
                else:
                    counts[i].append(lap_count)
            dates.append(record.date)
        if not counts:
            return False
        plt.clf()
        set_up_dates_xaxis(dates)
        # Ensure y-axis integer ticks only (makes sense - discrete data).
        # Max 10 coins, max 3 mushrooms
        plt.yticks(range(4) if field == "mushrooms" else range(11))
        for n, lap_counts in enumerate(counts, 1):
            plt.plot(dates, lap_counts, label=f"Lap {n}")
        plt.title(f"{field.capitalize()} against Date")
        plt.xlabel("Date")
        plt.ylabel(field.capitalize())
        plt.legend(loc="upper left")
        return True
    
    def plot_coins_against_date(self) -> bool:
        """Plots the coins against date graph."""
        return self._plot_item_against_date("coins")

    def plot_mushrooms_against_date(self) -> bool:
        """Plots the mushrooms against date graph."""
        return self._plot_item_against_date("mushrooms")


class GraphFrame(tk.Frame):
    """
    Stores a matplotlib image graph with the option to open interactive mode.
    """

    def __init__(
        self, master: GraphsToplevel, plot_command: Callable
    ) -> None:
        super().__init__(master)
        self.plot_command = plot_command
        success = self.plot_command()
        if not success:
            self.label = tk.Label(self, font=lato(25), text="No data")
            self.label.pack()
            self.graph_image = None
            return
        with io.BytesIO() as f:
            plt.savefig(f, format="png", dpi=GRAPH_IMAGE_DPI)
            f.seek(0)
            self.graph_image = tk.PhotoImage(data=f.read(), format="png")
        self.graph_image_label = tk.Label(self, image=self.graph_image)
        self.view_button = ttk.Button(
            self, text="View", command=self.view)
        self.graph_image_label.pack(padx=5, pady=5)
        self.view_button.pack(padx=5, pady=5)

    def view(self) -> None:
        """Opens interactive window for particular graph."""
        self.plot_command()
        plt.show(block=False)


class ExportationToplevel(tk.Toplevel):
    """Exportation to CSV, XLSX, DOCX and PDF."""

    def __init__(self, master: CourseCcAnalysis) -> None:
        super().__init__()
        self.master: CourseCcAnalysis = master
        self.title(f"{master.title.cget('text')} - Data Exportation")
        self.grab_set()
        self.notebook = ttk.Notebook(self)
        self.table_exportation_frame = CourseCcTableExportationFrame(self)
        self.document_exportation_frame = DocumentExportationFrame(self)
        self.notebook.add(self.table_exportation_frame, text="CSV/XLSX")
        self.notebook.add(self.document_exportation_frame, text="DOCX/PDF")
        self.notebook.pack()


class CourseCcTableExportationFrame(TableExportationFrame):
    """Subclass for Course/CC table exportation."""

    def __init__(self, master: CourseCcAnalysis) -> None:
        super().__init__(master, COURSE_CC_EXPORT_TABLES)

    def get_columns_and_records(
        self, table: ExportTable
    ) -> tuple[tuple[str], list[tuple]]:
        """Returns the columns and records for a given table."""
        if table == ExportTable.raw_records:
            laps = get_lap_count(self.master.master.course)
            columns = get_raw_records_columns(laps)
            records = []
            for record in self.master.master.records:
                export_record = [
                    record.course, record.is200, record.date, record.time,
                    record.player, record.country, record.days]
                for field in ("lap_times", "coins", "mushrooms"):
                    if getattr(record, field) is None:
                        export_record.extend([None] * laps)
                    else:
                        export_record.extend(getattr(record, field))
                export_record.extend(
                    (record.character, record.kart, record.tyres,
                        record.glider, record.video_link))
                records.append(export_record)
            return columns, records
        columns = EXPORT_TABLES_COLUMNS[table]
        tables = self.master.master.tables_frame
        records = get_records_from_tables(tables, table)
        return columns, records


class DocumentExportationFrame(tk.Frame):
    """Exportation to DOCX or PDF."""

    def __init__(self, master: ExportationToplevel) -> None:
        super().__init__(master)
        self.title = tk.Label(
            self, font=lato(25, True), text="Export Document")
        self.info_label = tk.Label(
            self, text="Select the components below to generate.")
        self.basic_label = tk.Label(self, font=lato(15, True), text="Basic")
        self.basic_frame = BasicDocumentExportationOptions(self)
        self.tables_label = tk.Label(self, font=lato(15, True), text="Tables")
        self.tables_frame = TableDocumentExportationOptions(
            self, DOCUMENT_COURSE_CC_EXPORT_TABLES)
        self.graphs_label = tk.Label(self, font=lato(15, True), text="Graphs")
        self.graphs_frame = GraphDocumentExportationFrame(self)
        self.export_docx_button = ttk.Button(
            self, text="Export DOCX", command=self.export_docx)
        self.export_pdf_button = ttk.Button(
            self, text="Export PDF [WARNING: Unstable!]",
            command=self.export_pdf)
        
        self.title.pack(pady=5)
        self.info_label.pack(pady=5)
        self.basic_label.pack(pady=5)
        self.basic_frame.pack(pady=5)
        self.tables_label.pack(pady=5)
        self.tables_frame.pack(pady=5)
        self.graphs_label.pack(pady=5)
        self.graphs_frame.pack(pady=5)
        self.export_docx_button.pack(pady=5)
        self.export_pdf_button.pack(pady=5)
    
    def _generate_docx(self) -> docx.document.Document:
        # Generate DOCX from the configuration options.
        basic_options = self.basic_frame.options
        tables = self.tables_frame.selected
        rows_per_table = self.tables_frame.max_count
        graphs = self.graphs_frame.selected
        document: docx.document.Document = docx.Document()
        document.add_heading(self.master.master.title.cget("text"), 0)
        if basic_options.current_records or basic_options.summary_stats:
            self._add_basic(document, basic_options)
        if tables:
            self._add_tables(document, tables, rows_per_table)
        if graphs:
            self._add_graphs(document, graphs)
        document.add_heading("Data Context", 1)
        options_frame = self.master.master.options_frame
        date_range_text = options_frame.date_range_label.cget("text")
        data_time_text = options_frame.data_time_label.cget("text")
        document.add_paragraph(date_range_text)
        document.add_paragraph(data_time_text)
        return document
    
    def _add_basic(
        self, document: docx.document.Document,
        basic_options: BasicDocumentOptions
    ) -> None:
        # Adds basic info to the document.
        if basic_options.current_records:
            current_records_frame = self.master.master.current_records_frame
            current_records = current_records_frame.current_records
            document.add_heading(
                f"Current Record{'s' if len(current_records) > 1 else ''}", 1)
            if not current_records:
                document.add_paragraph("No current record.")
            else:
                records_parts = []
                for record in current_records:
                    finish_time = ms_to_finish_time(record.time)
                    lap_times = record.lap_times
                    if lap_times is not None:
                        lap_times = tuple(
                            round(lap / 1000, 3) for lap in lap_times)
                    parts = (
                        f"Date: {record.date}", f"Time: {finish_time}",
                        f"Player: {record.player}",
                        f"Country: {record.country}",
                        f"Days lasted: {record.days}",
                        f"Lap times: {lap_times}",
                        f"Coins per lap: {record.coins}",
                        f"Mushrooms per lap: {record.mushrooms}",
                        f"Character: {record.character}",
                        f"Kart: {record.kart}", f"Tyres: {record.tyres}",
                        f"Glider: {record.glider}",
                        f"Video Link: {record.video_link}")
                    records_parts.append(parts)
                if len(records_parts) == 1:
                    for part in records_parts[0]:
                        document.add_paragraph(part, style="List Bullet")
                else:
                    for parts in records_parts:
                        document.add_paragraph(
                            " | ".join(parts), style="List Bullet")
        if basic_options.summary_stats:
            document.add_heading("Summary Stats", 1)
            info_label = self.master.master.info_label
            points = map(str.strip, info_label.cget("text").split("|"))
            for point in points:
                document.add_paragraph(point, style="List Bullet")

    def _add_tables(
        self, document: docx.document.Document,
        tables: list[ExportTable], rows_per_table: int
    ) -> None:
        # Adds tables to the document.
        document.add_heading("Tables", 1)
        tables_frame = self.master.master.tables_frame
        for table in tables:
            document.add_heading(table.value, 2)
            columns = ("#", *EXPORT_TABLES_COLUMNS[table])
            records = get_records_from_tables(
                tables_frame, table)[:rows_per_table]
            add_doctable(document, columns, records)
    
    def _add_graphs(
        self, document: docx.document.Document, graphs: list[ExportGraph]
    ) -> None:
        # Adds graphs to the document.
        document.add_heading("Graphs", 1)
        graph_functions = {
            ExportGraph.times: GraphsToplevel.plot_time_against_date,
            ExportGraph.lap_times: GraphsToplevel.plot_laps_against_date,
            ExportGraph.coins: lambda m: (
                GraphsToplevel._plot_item_against_date(m, "coins")),
            ExportGraph.mushrooms: lambda m: (
                GraphsToplevel._plot_item_against_date(m, "mushrooms"))
        }
        for graph in graphs:
            graph_function = graph_functions[graph]
            # Duck typing - keep it simple and dirty.
            # The master of this widget has master with records as needed.
            success = graph_function(self.master)
            if not success:
                document.add_paragraph(f"{graph.value} graph unavailable.")
                continue
            with io.BytesIO() as graph_bytes:
                plt.savefig(graph_bytes, format="png", dpi=GRAPH_IMAGE_DPI)
                graph_bytes.seek(0)
                document.add_picture(graph_bytes, GRAPH_WIDTH, GRAPH_HEIGHT)

    def export_docx(self) -> None:
        """Exports to DOCX."""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx", filetypes=(("DOCX", ".docx"),),
            title="Save DOCX", parent=self)
        if not file_path:
            return
        try:
            document = self._generate_docx()
            document.save(file_path)
        except Exception as e:
            messagebox.showerror(
                "Error",
                    "Unfortunately, an error occurred "
                    f"while exporting to DOCX: {e}", parent=self)
    
    def export_pdf(self) -> None:
        """
        Exports to PDF by first creating DOCX, then converting to PDF.
        Not the most stable of operations but at least it's possible.
        """
        file_path = filedialog.asksaveasfilename(
            defaultextension=".pdf", filetypes=(("PDF", ".pdf"),),
            title="Save PDF", parent=self)
        if not file_path:
            return
        try:
            document = self._generate_docx()
            save_docx_as_pdf(document, file_path)
        except Exception as e:
            messagebox.showerror(
                "Error",
                    "Unfortunately, an error occurred "
                    f"while exporting to PDF: {e}", parent=self)


class BasicDocumentExportationOptions(tk.Frame):
    """
    Sets whether or not to include current records
    and summary stats in the document.
    """

    def __init__(self, master: DocumentExportationFrame) -> None:
        super().__init__(master)
        self._current_records = tk.BooleanVar(value=True)
        self._summary_stats = tk.BooleanVar(value=True)
        self.current_records_checkbutton = ttk.Checkbutton(
            self, text="Current Records", variable=self._current_records)
        self.summary_stats_checkbutton = ttk.Checkbutton(
            self, text="Summary Stats", variable=self._summary_stats)
        self.current_records_checkbutton.pack()
        self.summary_stats_checkbutton.pack()
    
    @property
    def options(self) -> BasicDocumentOptions:
        return BasicDocumentOptions(
            self._current_records.get(), self._summary_stats.get())


class GraphDocumentExportationFrame(tk.Frame):
    """Allows the graphs to generate into the document to be selected."""

    def __init__(self, master: DocumentExportationFrame) -> None:
        super().__init__(master)
        self._selected = [
            tk.BooleanVar(value=True)
            for _ in range(len(COURSE_CC_EXPORT_TABLES))]
        for var, table in zip(self._selected, COURSE_CC_EXPORT_GRAPHS):
            checkbutton = ttk.Checkbutton(self, text=table.value, variable=var)
            checkbutton.pack(padx=5)
        
    @property
    def selected(self) -> list[ExportGraph]:
        return [
            table for var, table
                in zip(self._selected, COURSE_CC_EXPORT_GRAPHS) if var.get()]
