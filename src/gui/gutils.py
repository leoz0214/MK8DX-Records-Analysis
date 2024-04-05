"""GUI-specific utilities (avoid module name clash with main utilities)."""
import csv
import datetime as dt
import pathlib
import platform
import tempfile
import tkinter as tk
from collections import Counter
from contextlib import suppress
from tkinter import filedialog
from tkinter import messagebox
from tkinter import ttk
from typing import Callable, Union

try:
    import comtypes.client as comtypesclient
except ImportError:
    comtypesclient = None
import docx
with suppress(ImportError):
    import docx2pdf
import matplotlib.dates as mdates
import openpyxl
import openpyxl.styles
import pyglet
from tkcalendar import Calendar

import course_cc
import general
from const import LATO_FONT_FILE
from utils import Record, get_lap_count


pyglet.font.add_file(str(LATO_FONT_FILE))

DAYS_BY_PLAYER_COLS = ("Player", "Days Held", "%")
RECORDS_BY_PLAYER_COLS = ("Player", "Records", "%")
DAYS_BY_COUNTRY_COLS = ("Country", "Days Held", "%")
RECORDS_BY_COUNTRY_COLS = ("Country", "Records", "%")
RECORDS_BY_CHARACTER_COLS = ("Character", "Records", "%")
RECORDS_BY_KART_COLS = ("Kart", "Records", "%")
RECORDS_BY_TYRES_COLS = ("Tyres", "Records", "%")
RECORDS_BY_GLIDER_COLS = ("Glider", "Records", "%")


DEFAULT_RANK_TABLE_HEIGHT = 3
RANK_COL_WIDTH = 40
PLAYER_COL_WIDTH = 200
COUNTRY_COL_WIDTH = 200
COURSE_COL_WIDTH = 250
CC_COL_WIDTH = 60
IMPROVEMENT_COL_WIDTH = 125
BUILD_COL_WIDTH = 200
COUNT_COL_WIDTH = 100
PERCENTAGE_COL_WIDTH = 70
# Special quick dates that may be referenced.
SPECIAL_DATES = {
    dt.date(2023, 3, 8): "Day before Patch 1",
    dt.date(2023, 3, 9): "Day of Patch 1",
    dt.date(2023, 7, 11): "Day before Patch 2",
    dt.date(2023, 7, 12): "Day of Patch 2"
}
MIN_DATE = dt.date(2017, 1, 1)

MIN_DOCUMENT_TABLE_ROWS = 1
DEFAULT_DOCUMENT_TABLE_ROWS = 3
MAX_DOCUMENT_TABLE_ROWS = 10

WORD_EXPORT_PDF_ID = 17


def lato(size: int, bold: bool = False, italic: bool = False) -> tuple:
    """Utility function for the Lato font."""
    font = ("Lato", size)
    if bold:
        font += ("bold",)
    if italic:
        font += ("italic",)
    return font


def bool_to_state(expression: bool) -> str:
    """Returns 'normal' if True, 'disabled' if False."""
    return "normal" if expression else "disabled"


def get_date_range_string(
    start_date: dt.date | None, end_date: dt.date | None
) -> str:
    """Returns a string representing a date range."""
    if start_date is None and end_date is None:
        return "All Time"
    elif start_date is None:
        return f"Start - {end_date}"
    elif end_date is None:
        return f"{start_date} - Current"
    return f"{start_date} to {end_date}"


def sort_by_days_held(
    records: list[Record], uniques: set[str], field: str
) -> list[tuple]:
    """Returns records ranking days by player or country."""
    if not uniques:
        return []
    days_by = dict.fromkeys(uniques, 0)
    total_days = 0
    for record in records:
        days_by[getattr(record, field)] += record.days
        total_days += record.days
    # Sorts from most days held to least.
    days_by = dict(
        sorted(days_by.items(), key=lambda item: item[1], reverse=True))
    days_by_records = [
        (key, days, round(days / total_days * 100, 2))
        for key, days in days_by.items()]
    return days_by_records


def sort_by_records_count(records: list[Record], field: str) -> list[tuple]:
    """
    Returns records ranking number of records by player, country,
    character etc.
    """
    records_counts = Counter(getattr(record, field) for record in records)
    if None in records_counts:
        # Ignore empty values.
        records_counts.pop(None)
    if not records:
        # Somehow not a single record has the relevant data - no return.
        return []
    # Sorts from most records to least.
    records_counts = dict(
        sorted(records_counts.items(), key=lambda item: item[1], reverse=True))
    records_count_records = [
        (key, count, round(count / len(records) * 100, 2))
        for key, count in records_counts.items()]
    return records_count_records


def get_raw_records_columns(laps: int) -> tuple[str]:
    """
    Returns columns for raw records (where lap count determines
    number of columns for lap times, coins, mushrooms).
    """
    return (
        "Course", "Is200", "Date", "Time", "Player", "Country", "Days",
        *(f"Lap {n}" for n in range(1, laps + 1)),
        *(f"Coins {n}" for n in range(1, laps + 1)),
        *(f"Mushrooms {n}" for n in range(1, laps + 1)),
        "Character", "Kart", "Tyres", "Glider", "Video Link")


class RankTable(tk.Frame):
    """
    Convenient rank stat table which stores a fixed treeview
    determined at the point of creation. Can be used to store
    a table of data (top countries, top players etc...).
    """

    def __init__(
        self, master: tk.Widget,
        label: str, columns: tuple[str], rows: list[tuple],
        height: int = DEFAULT_RANK_TABLE_HEIGHT,
        column_widths: tuple[int] = None
    ) -> None:
        super().__init__(master)
        self.label = tk.Label(self, font=lato(15, True), text=label)
        columns = ("#", *columns)
        ttk.Style().configure("Treeview", rowheight=30)
        self.treeview = ttk.Treeview(
            self, columns=columns, show="headings", height=height)
        for column in columns:
            self.treeview.heading(column, text=column)
        if column_widths is not None:
            column_widths = (RANK_COL_WIDTH, *column_widths)
            for column, width in zip(columns, column_widths):
                self.treeview.column(column, width=width)
        for n, row in enumerate(rows, 1):
            row = (n, *row)
            self.treeview.insert("", "end", values=row)
        self.scrollbar = ttk.Scrollbar(
            self, orient="vertical", command=self.treeview.yview)
        self.treeview.config(yscrollcommand=self.scrollbar.set)
        self.label.grid(row=0, column=0, columnspan=2)
        self.treeview.grid(row=1, column=0)
        self.scrollbar.grid(row=1, column=1, sticky="ns")


class TablesFrame(tk.Frame):
    """Various tables providing various rank stats."""

    def __init__(
        self, master: Union[
            "course_cc.CourseCcAnalysis", "general.OverallRecordsDataFrame",
            "general.CurrentRecordsDataFrame"]
        ) -> None:
        super().__init__(master)
        self.days_by_player_records = sort_by_days_held(
            master.records, master.unique_players, "player")
        self.days_by_player_frame = RankTable(
            self, "Days by Player", DAYS_BY_PLAYER_COLS,
            self.days_by_player_records, column_widths=(
                PLAYER_COL_WIDTH, COUNT_COL_WIDTH, PERCENTAGE_COL_WIDTH))
        self.player_record_counts = (
            sort_by_records_count(master.records, "player"))
        self.player_records_counts_frame = RankTable(
            self, "Records by Player", RECORDS_BY_CHARACTER_COLS,
            self.player_record_counts, column_widths=(
                PLAYER_COL_WIDTH, COUNT_COL_WIDTH, PERCENTAGE_COL_WIDTH))
        self.days_by_country_records = sort_by_days_held(
            master.records, master.unique_countries, "country")
        self.days_by_country_frame = RankTable(
            self, "Days by Country", DAYS_BY_COUNTRY_COLS,
            self.days_by_country_records, column_widths=(
                COUNTRY_COL_WIDTH, COUNT_COL_WIDTH, PERCENTAGE_COL_WIDTH))
        self.country_record_counts = (
            sort_by_records_count(master.records, "country"))
        self.country_records_counts_frame = RankTable(
            self, "Records by Country", RECORDS_BY_COUNTRY_COLS,
            self.country_record_counts, column_widths=(
                COUNTRY_COL_WIDTH, COUNT_COL_WIDTH, PERCENTAGE_COL_WIDTH))

        self.character_record_counts = (
            sort_by_records_count(master.records, "character"))
        self.character_records_counts_frame = RankTable(
            self, "Records by Character", RECORDS_BY_CHARACTER_COLS,
            self.character_record_counts, column_widths=(
                BUILD_COL_WIDTH, COUNT_COL_WIDTH, PERCENTAGE_COL_WIDTH))
        self.kart_record_counts = sort_by_records_count(master.records, "kart")
        self.kart_records_counts_frame = RankTable(
            self, "Records by Kart", RECORDS_BY_KART_COLS,
            self.kart_record_counts, column_widths=(
                BUILD_COL_WIDTH, COUNT_COL_WIDTH, PERCENTAGE_COL_WIDTH))
        self.tyres_record_counts = (
            sort_by_records_count(master.records, "tyres"))
        self.tyres_records_counts_frame = RankTable(
            self, "Records by Tyres", RECORDS_BY_TYRES_COLS,
            self.tyres_record_counts, column_widths=(
                BUILD_COL_WIDTH, COUNT_COL_WIDTH, PERCENTAGE_COL_WIDTH))
        self.glider_record_counts = (
            sort_by_records_count(master.records, "glider"))
        self.glider_records_counts_frame = RankTable(
            self, "Records by Glider", RECORDS_BY_GLIDER_COLS,
            self.glider_record_counts, column_widths=(
                BUILD_COL_WIDTH, COUNT_COL_WIDTH, PERCENTAGE_COL_WIDTH))
        self.days_by_player_frame.grid(row=0, column=0, padx=3, pady=3)
        self.player_records_counts_frame.grid(row=0, column=1, padx=3, pady=3)
        self.days_by_country_frame.grid(row=0, column=2, padx=3, pady=3)
        self.country_records_counts_frame.grid(row=0, column=3, padx=3, pady=3)
        self.character_records_counts_frame.grid(
            row=1, column=0, padx=3, pady=3)
        self.kart_records_counts_frame.grid(row=1, column=1, padx=3, pady=3)
        self.tyres_records_counts_frame.grid(row=1, column=2, padx=3, pady=3)
        self.glider_records_counts_frame.grid(row=1, column=3, padx=3, pady=3)


class MillisecondsFormatter(mdates.DateFormatter):
    """
    Monkey patches the matplotlib date formatter to allow for milliseconds.
    Note: assumes last part of the string is indeed microseconds.
    Resolution of lap/finish times is indeed milliseconds only
    (no truncation issues).
    """

    def __call__(self, x, pos=0):
        """
        Monkey patched exactly from source except strips last 3 digits
        to show ms instead of us: may break in future... be warned!
        """
        result = mdates.num2date(x, self.tz).strftime(self.fmt)[:-3]
        return mdates._wrap_in_tex(result) if self._usetex else result


class DatePicker(tk.Toplevel):
    """Allows a date to be selected using the calendar."""

    def __init__(self, title: str, callback: Callable) -> None:
        super().__init__()
        self.title(title)
        self._initial_grab = self.grab_current()
        self.callback = callback
        self.protocol("WM_DELETE_WINDOW", self.close)
        self.grab_set()
        self.title_label = tk.Label(
            self, font=lato(25, True), text=title)
        max_date = dt.date.today() + dt.timedelta(days=7)
        self.calendar = Calendar(self, mindate=MIN_DATE, maxdate=max_date)
        self.selected_date_label = tk.Label(
            self, text=f"Selected: {self.selected}")
        self.calendar.bind(
            "<<CalendarSelected>>", lambda *_: self.update_selected_date())
        self.select_button = ttk.Button(
            self, text="Select", command=self.select)
        self.quick_dates_label = tk.Label(self, text="Quick dates:")

        self.title_label.pack(padx=10, pady=10)
        self.calendar.pack(padx=10, pady=10)
        self.selected_date_label.pack(padx=10, pady=10)
        self.select_button.pack(padx=10, pady=10)
        self.quick_dates_label.pack(padx=10, pady=10)

        for date, text in SPECIAL_DATES.items():
            quick_date_button = ttk.Button(
                self, text=f"{date} - {text}", width=30,
                command=self.select_quick_date(date))
            quick_date_button.pack(padx=10, pady=5)
        
    @property
    def selected(self) -> dt.date:
        return self.calendar._sel_date

    def update_selected_date(self) -> None:
        """Updates the selected date label upon change."""
        self.selected_date_label.config(text=f"Selected: {self.selected}")

    def select(self) -> None:
        """Selects the current date in the calendar."""
        self.close()
        self.callback(self.selected)
    
    def select_quick_date(self, date: dt.date) -> Calendar:
        """Wrapper for quick date selection."""
        def wrapper() -> None:
            self.close()
            self.callback(date)
        return wrapper

    def close(self) -> None:
        """Closes the date picker, returning control to master."""
        self.destroy()
        self._initial_grab.grab_set()


class DateRangeSelect(tk.Frame):
    """Allows the user to select a date range for analysis."""

    def __init__(
        self, master: tk.Widget,
        initial_start_date: dt.date | None = None,
        initial_end_date: dt.date | None = None
    ) -> None:
        super().__init__(master)
        self.start_date = initial_start_date
        self.end_date = initial_end_date
        self.start_date_button = ttk.Button(
            self, text="Set Start Date",
            command=lambda: DatePicker("Set Start Date", self.set_start_date))
        self.end_date_button = ttk.Button(
            self, text="Set End Date",
            command=lambda: DatePicker("Set End Date", self.set_end_date))
        self.reset_start_date_button = ttk.Button(
            self, text="Reset Start Date",
            command=lambda: self.set_start_date(None))
        self.reset_end_date_button = ttk.Button(
            self, text="Reset End Date",
            command=lambda: self.set_end_date(None))
        self.label = tk.Label(self)
        self.update_label()
        self.label.grid(row=0, column=0, padx=5)
        self.start_date_button.grid(row=0, column=1, padx=5)
        self.end_date_button.grid(row=0, column=2, padx=5)
        self.reset_start_date_button.grid(row=0, column=3, padx=5)
        self.reset_end_date_button.grid(row=0, column=4, padx=5)
    
    def set_start_date(self, date: dt.date | None) -> None:
        """Sets the start date."""
        if (
            date is not None and self.end_date is not None
            and date > self.end_date
        ):
            messagebox.showerror(
                "Error", "Start date must not be later than end date.",
                parent=self)
            return
        self.start_date = date
        self.update_label()
    
    def set_end_date(self, date: dt.date | None) -> None:
        """Sets the end date."""
        if (
            date is not None and self.start_date is not None
            and date < self.start_date
        ):
            messagebox.showerror(
                "Error", "End date must not be earlier than start date.",
                parent=self)
            return
        self.end_date = date
        self.update_label()
    
    def update_label(self) -> None:
        """Updates the label based on start and end date."""
        date_range = get_date_range_string(self.start_date, self.end_date)
        self.label.config(text=f"Date Range: {date_range}")


class UpdateDateRangeToplevel(tk.Toplevel):
    """Allows the date range to be updated even from an initialised tab"""

    def __init__(
        self, callback: Callable,
        current_start_date: dt.date | None, current_end_date: dt.date | None
    ) -> None:
        super().__init__()
        self.title("Update Date Range")
        self.callback = callback
        self.grab_set()
        self.title_label = tk.Label(
            self, font=lato(25, True), text="Update Date Range")
        self.date_selection_frame = DateRangeSelect(
            self, current_start_date, current_end_date)
        self.update_button = ttk.Button(
            self, text="Update", command=self.update_date_range)
        self.title_label.pack(padx=5, pady=5)
        self.date_selection_frame.pack(padx=5, pady=5)
        self.update_button.pack(padx=5, pady=5)
    
    def update_date_range(self) -> None:
        """Updates the date range as required."""
        self.destroy()
        start_date = self.date_selection_frame.start_date
        end_date = self.date_selection_frame.end_date
        self.callback(start_date, end_date)


class TableExportationFrame(tk.Frame):
    """Exportation to CSV or XLSX."""

    def __init__(
        self, master: Union[
            "course_cc.ExportationToplevel", "general.ExportationToplevel"],
        tables: tuple
    ) -> None:
        super().__init__(master)
        self.title = tk.Label(self, font=lato(25, True), text="Export Tables")
        self.info_label = tk.Label(
            self, text=(
                "Select one or more tables below to export.\n"
                "Note: CSV is limited to one table, XLSX unlimited."))
        self.options_frame = TableExportationOptions(self, tables)
        self.export_csv_button = ttk.Button(
            self, text="Export CSV", command=self.export_csv)
        self.export_xlsx_button = ttk.Button(
            self, text="Export XLSX", command=self.export_xlsx)
        self.update_button_states()
        self.title.pack(pady=5)
        self.info_label.pack(pady=5)
        self.options_frame.pack(pady=5)
        self.export_csv_button.pack(pady=5)
        self.export_xlsx_button.pack(pady=5)
    
    def update_button_states(self) -> None:
        """Updates the export button states."""
        selected_tables = self.options_frame.tables
        # Exactly one table selected for CSV export to be allowed.
        self.export_csv_button.config(
            state=bool_to_state(len(selected_tables) == 1))
        # One or more tables selected for XLSX export to be allowed.
        self.export_xlsx_button.config(state=bool_to_state(selected_tables))
    
    def export_csv(self) -> None:
        """Exports a single table to CSV."""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=(("CSV", ".csv"),),
            title="Save CSV", parent=self)
        if not file_path:
            return
        table = self.options_frame.tables[0]
        try:
            columns, records = self.get_columns_and_records(table)
            with open(file_path, "w", encoding="utf8") as f:
                writer = csv.writer(f, lineterminator="\n")
                writer.writerow(columns)
                writer.writerows(records)
        except Exception as e:
            messagebox.showerror(
                "Error",
                    "Unfortunately, an error occurred "
                    f"while saving the CSV: {e}", master=self)

    def export_xlsx(self) -> None:
        """Exports one or more tables to XLSX (multiple sheets)."""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx", filetypes=(("XLSX", ".xlsx"),),
            title="Save XLSX", parent=self)
        if not file_path:
            return
        try:
            workbook = openpyxl.Workbook()
            workbook.remove(workbook["Sheet"])
            tables = self.options_frame.tables
            for table in tables:
                columns, records = self.get_columns_and_records(table)
                sheet = workbook.create_sheet(table.value)
                for i, column in enumerate(columns, 1):
                    cell = sheet.cell(row=1, column=i)
                    cell.value = column
                    cell.font = openpyxl.styles.Font(bold=True)
                for record in records:
                    sheet.append(record)
            workbook.save(file_path)
        except Exception as e:
            messagebox.showerror(
                "Error",
                    "Unfortunately, an error occurred "
                    f"while saving the XLSX: {e}", master=self)
    
    def get_columns_and_records(self, _) -> tuple[tuple[str], list[tuple]]:
        # Needs to be subclassed and implemented in derived class.
        raise NotImplementedError


class TableExportationOptions(tk.Frame):
    """Contains the checkbuttons to select the tables to export."""

    def __init__(self, master: TableExportationFrame, tables: tuple) -> None:
        super().__init__(master)
        self._tables = tables
        self._selected = [
            tk.BooleanVar(value=False)
            for _ in range(len(tables))]
        for var in self._selected:
            var.trace("w", lambda *_: master.update_button_states())
        for i, table in enumerate(self._tables):
            checkbutton = ttk.Checkbutton(
                self, text=table.value, variable=self._selected[i])
            checkbutton.pack()
    
    @property
    def tables(self) -> list:
        return [
            table for var, table in zip(self._selected, self._tables)
                if var.get()]


class TableDocumentExportationOptions(tk.Frame):
    """Allows the tables to generate into the document to be selected."""

    def __init__(self, master: tk.Frame, tables: tuple) -> None:
        super().__init__(master)
        self.tables = tables
        self._selected = [
            tk.BooleanVar(value=True)
            for _ in range(len(tables))]
        self._max_count = tk.IntVar(value=DEFAULT_DOCUMENT_TABLE_ROWS)
        for i, (var, table) in enumerate(
            zip(self._selected, self.tables)
        ):
            checkbutton = ttk.Checkbutton(self, text=table.value, variable=var)
            checkbutton.grid(row=i//2, column=i%2, padx=5)
        self.max_count_label = tk.Label(self, text="Display top")
        self.max_count_scale = tk.Scale(
            self, from_=MIN_DOCUMENT_TABLE_ROWS, to=MAX_DOCUMENT_TABLE_ROWS,
            orient="horizontal", variable=self._max_count, length=200)
        self.max_count_label.grid(
            row=10, column=0, padx=5, pady=5, sticky="ne")
        self.max_count_scale.grid(
            row=10, column=1, padx=5, pady=5, sticky="w")
    
    @property
    def max_count(self) -> int:
        # Top N only for each table (full tables too big in document).
        return self._max_count.get()
    
    @property
    def selected(self) -> list:
        return [
            table for var, table in zip(self._selected, self.tables)
                if var.get()]


def add_doctable(
    document: docx.document.Document, columns: tuple[str], records: list[tuple]
) -> None:
    """Adds a document table given columns and records."""
    doctable = document.add_table(rows=len(records) + 1, cols=len(columns))
    doctable.style = "Table Grid"
    doctable.autofit = True
    for i, column in enumerate(columns):
        cell = doctable.cell(0, i)
        cell.text = str(column)
        # Make cell (column) text bold.
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for i, record in enumerate(records, 1):
        record = (i, *record)
        for j, value in enumerate(record):
            doctable.cell(i, j).text = str(value)
    for cell in doctable.columns[0].cells:
        cell.width = 50


def save_docx_as_pdf(document: docx.document.Document, file_path: str) -> None:
    """Attempts to save document to given file path as PDF."""
    # Create temp file to use as dummy DOCX.
    with tempfile.NamedTemporaryFile(
        "wb", suffix=".docx", delete=False
    ) as f:
        temp_file_path = pathlib.Path(f.name)
        document.save(f)
    try:
        if platform.system() == "Windows" and comtypesclient is not None:
            # Use MS Word on windows to attempt conversion.
            # Avoids weird behaviour of docx2pdf.
            word = None
            doc = None
            try:
                word = comtypesclient.CreateObject("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(str(temp_file_path))
                doc.SaveAs(file_path, FileFormat=WORD_EXPORT_PDF_ID)
            finally:
                if doc is not None:
                    doc.Close()
                if word is not None:
                    word.Quit()
        else:
            # For MacOS, attempt conversion using docx2pdf.
            docx2pdf.convert(temp_file_path, file_path, keep_active=True)
    finally:
        temp_file_path.unlink(missing_ok=True)